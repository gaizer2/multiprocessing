import multiprocessing
import time
import random
import docx



def passw(s):
    #Функция генерации паролей
    start = time.time()
    mass = "ptD23OTEgxhaEKiaPJTz"
    password = []

    for x in range(s):
        password.append(mass[random.randint(0, len(mass) - 1)])
    result = ''.join(password)
    print("Готовый пароль ", result)
    end = time.time()
    print("Время выполнения", end - start)
    print("============================PROCCES 1 CLOSE==========================")


def worker_function2(process_num):
    # Функция подсчета повторяющихся предложений
    start = time.time()
    document1 = docx.Document('C://Users/Gaizer/Desktop/1.docx')
    document2 = docx.Document('C://Users/Gaizer/Desktop/2.docx')
    array1 = []
    array2 = []
    g = 0
    for x in document1.paragraphs:
        array1 += x.text.split('.')

    array1 = list(filter(lambda x: x != "", array1))
    print("Предложений в файле 1: ", len(array1))

    for x in document2.paragraphs:
        array2 += x.text.split('.')
    array2 = list(filter(lambda x: x != "", array2))
    print("Предложений в файле 2: ", len(array2))

    for x in range(len(array1)):
        for y in range(len(array2)):
            if array2[y] == array1[x]:
                g += 1
    print("КОличество одинаковых предложений ", g, g / len(array1) * 100, "%")
    end = time.time()
    print("Время выполнения", end - start)
    print("============================PROCCES 2 CLOSE==========================")


def worker_function3(process_num):
    #Функция подсчета повторяющихся слов
    start = time.time()
    document1 = docx.Document('C://Users/Gaizer/Desktop/1.docx')
    document2 = docx.Document('C://Users/Gaizer/Desktop/2.docx')
    array1 = []
    array2 = []
    slova1 = []
    slova2 = []
    slova3 = []
    count = []
    g = 0
    for x in document1.paragraphs:
        array1 += x.text.split('.')

    array1 = list(filter(lambda x: x != "", array1))
    for x in range(len(array1)):
        slova1 += array1[x].split(' ')
    print("Слов в файле 1: ", len(slova1))

    for x in document2.paragraphs:
        array2 += x.text.split('.')
    array2 = list(filter(lambda x: x != "", array2))
    for x in range(len(array2)):
        slova2 += array2[x].split(' ')
    print("Слов в файле 2: ", len(slova2))

    for x in range(len(slova1)):
        for y in range(len(slova2)):
            if slova2[y] == slova1[x]:
                if slova2[y] in slova3:
                    continue
                else:
                    slova3.append(slova2[y])
                    g += 1

    for x in range(len(slova3)):
        id = 0
        for y in range(len(slova1)):
            if slova3[x] == slova1[y]:
                id += 1
        count.append(id)

    print("Количество одинаковых слов ", g, g / len(slova1) * 100, "%")
    end = time.time()
    print("Время выполнения", end - start)
    print("============================PROCCES 3 CLOSE==========================")


if __name__ == "__main__":
    s = int(input("Введите длинну пароля "))
    start = time.time()
    # Создание процессов
    process1 = multiprocessing.Process(target=passw, args=(s,))
    process2 = multiprocessing.Process(target=worker_function2, args=(2,))
    process3 = multiprocessing.Process(target=worker_function3, args=(2,))
    # Запуск процессов
    process1.start()
    process2.start()
    process3.start()
    # Завершение процессов
    process1.join()
    process2.join()
    process3.join()
    end = time.time()
    print("Время выполнения", end - start)
